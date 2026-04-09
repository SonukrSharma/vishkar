import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.models.blueprint import BlueprintResult


OUTPUT_DIR = "generated_blueprints"

# Brand colors (hex strings for XML, RGBColor for python-docx)
PRIMARY     = "1e3a5f"
ACCENT      = "2e86ab"
GRAY        = "647487"
CODE_BG     = "f4f4f4"
CALLOUT_BG  = "e8f4fd"
ROW_ALT     = "f0f6fb"


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _shade_paragraph(p, fill: str):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _set_cell_bg(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _add_paragraph_border(p, side: str, color: str, sz: str = "6"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), "4")
    el.set(qn("w:color"), color)
    pBdr.append(el)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Inline markdown parser  (**bold**, *italic*, `code`)
# ---------------------------------------------------------------------------

def _add_inline(p, text: str, base_size: int = 10):
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    for part in pattern.split(text):
        if not part:
            continue
        run = p.add_run()
        run.font.size = Pt(base_size)
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
        else:
            run.text = part


# ---------------------------------------------------------------------------
# Block elements
# ---------------------------------------------------------------------------

def _add_code_block(doc: Document, code: str):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    for line in code.split("\n"):
        p = doc.add_paragraph()
        _shade_paragraph(p, CODE_BG)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2e)
    doc.add_paragraph().paragraph_format.space_before = Pt(2)


def _add_callout_box(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        _shade_paragraph(p, CALLOUT_BG)
        _add_paragraph_border(p, "left", ACCENT, "24")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _add_inline(p, line if line else " ", base_size=10)


def _add_markdown_table(doc: Document, lines: list[str]):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            text = row_data[c_idx] if c_idx < len(row_data) else ""

            if r_idx == 0:
                _set_cell_bg(cell, PRIMARY)
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            else:
                _set_cell_bg(cell, ROW_ALT if r_idx % 2 == 0 else "ffffff")
                _add_inline(p, text, base_size=9)

    doc.add_paragraph()


def _add_checkbox(doc: Document, text: str, checked: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    sym = p.add_run("☑  " if checked else "☐  ")
    sym.font.color.rgb = RGBColor(0x2e, 0x86, 0xab)
    sym.font.size = Pt(10)
    _add_inline(p, text)


# ---------------------------------------------------------------------------
# Content parser
# ---------------------------------------------------------------------------

def _parse_content(doc: Document, content: str):
    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        s = line.strip()

        # Code block
        if s.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        # Callout / blockquote
        if s.startswith("> "):
            callout = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                callout.append(lines[i].strip()[2:])
                i += 1
            _add_callout_box(doc, callout)
            continue

        # Table
        if s.startswith("|") and "|" in s[1:]:
            table_lines = []
            while i < len(lines):
                ts = lines[i].strip()
                if not ts.startswith("|"):
                    break
                # skip separator rows  |---|---|
                if not re.match(r"^\|[-| :]+\|$", ts):
                    table_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, table_lines)
            continue

        # Headings
        if s.startswith("#### "):
            doc.add_heading(s[5:], level=4)
        elif s.startswith("### "):
            p = doc.add_heading(s[4:], level=3)
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
        elif s.startswith("## "):
            p = doc.add_heading(s[3:], level=2)
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
        elif s.startswith("# "):
            p = doc.add_heading(s[2:], level=2)
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

        # Checkbox
        elif re.match(r"^- \[[ xX]\]", s):
            checked = s[3] in ("x", "X")
            _add_checkbox(doc, s[6:].strip(), checked)

        # Bullet
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, s[2:])

        # Numbered list
        elif re.match(r"^\d+\. ", s):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\. ", "", s))

        # Horizontal rule
        elif re.match(r"^[-*_]{3,}$", s):
            p = doc.add_paragraph()
            _add_paragraph_border(p, "bottom", "dddddd", "6")

        # Empty
        elif s == "":
            pass

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _add_inline(p, s)

        i += 1


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def _add_cover_page(doc: Document, result: BlueprintResult, answers_dict: dict):
    for _ in range(7):
        doc.add_paragraph()

    # App name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(result.app_name.upper())
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project Blueprint")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2e, 0x86, 0xab)

    doc.add_paragraph()

    # Divider
    p = doc.add_paragraph()
    _add_paragraph_border(p, "bottom", ACCENT, "12")

    doc.add_paragraph()

    # Generated by
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated by Vishkar")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x87)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x87)

    doc.add_paragraph()

    # Stack
    stack = (
        f"{answers_dict.get('tech_stack_backend', '')}  ·  "
        f"{answers_dict.get('tech_stack_frontend', '')}  ·  "
        f"{answers_dict.get('tech_stack_database', '')}"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(stack)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x87)

    doc.add_page_break()


def _add_toc(doc: Document, result: BlueprintResult):
    p = doc.add_heading("Table of Contents", level=1)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    doc.add_paragraph()

    for chapter in result.chapters:
        p = doc.add_paragraph()
        run = p.add_run(f"Chapter {chapter.chapter_number}   —   {chapter.title}")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()


def _add_chapter_header(doc: Document, number: int, title: str):
    # Chapter label
    p = doc.add_paragraph()
    run = p.add_run(f"CHAPTER {number}")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2e, 0x86, 0xab)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2)

    # Title
    p = doc.add_heading(title, level=1)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
        p.runs[0].font.size = Pt(22)
    _add_paragraph_border(p, "bottom", ACCENT, "8")

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_document(result: BlueprintResult, answers_dict: dict) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    _add_cover_page(doc, result, answers_dict)
    _add_toc(doc, result)

    for chapter in result.chapters:
        _add_chapter_header(doc, chapter.chapter_number, chapter.title)
        _parse_content(doc, chapter.content)
        doc.add_page_break()

    filename = (
        f"{result.app_name.replace(' ', '_')}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath
